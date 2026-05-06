"""tools/files/controller.py - File operations."""

import os
import re
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from config.logger import get_logger
from config.settings import USER_NAME, WORKSPACE_DIR

log = get_logger("files")
IS_WIN = sys.platform == "win32"

FOLDERS = {
    "downloads": Path.home() / "Downloads",
    "desktop": Path.home() / "Desktop",
    "documents": Path.home() / "Documents",
    "pictures": Path.home() / "Pictures",
    "music": Path.home() / "Music",
    "videos": Path.home() / "Videos",
    "workspace": WORKSPACE_DIR,
}

TYPE_MAP = {
    "Images": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg"],
    "Videos": [".mp4", ".mkv", ".avi", ".mov", ".wmv", ".flv"],
    "Documents": [".pdf", ".docx", ".doc", ".txt", ".xlsx", ".pptx", ".csv", ".md"],
    "Music": [".mp3", ".wav", ".flac", ".aac", ".ogg"],
    "Archives": [".zip", ".rar", ".7z", ".tar", ".gz"],
    "Code": [".py", ".js", ".html", ".css", ".java", ".cpp", ".ts"],
}


class FilesController:
    def __init__(self, output_dir: Path | None = None, memory=None):
        self.output_dir = Path(output_dir) if output_dir else FOLDERS["documents"]
        self.memory = memory
        self._recent_files = []

    def _resolve_path(self, path_or_query: str) -> Path:
        candidate = Path((path_or_query or "").strip().strip("\"'"))
        search_roots = list(FOLDERS.values()) + [Path.home(), Path.cwd()]

        if candidate.exists():
            return candidate

        if not candidate.is_absolute():
            for root in (Path.cwd(), WORKSPACE_DIR):
                joined = root / candidate
                if joined.exists():
                    return joined

        needle = candidate.name if candidate.name else str(candidate)
        for root in search_roots:
            if not root.exists():
                continue
            for found in root.rglob(f"*{needle}*"):
                return found
        return candidate

    def extract_text(self, path_or_query: str, max_chars: int = 4000) -> tuple[Path | None, str]:
        path = self._resolve_path(path_or_query)
        if not path.exists() or not path.is_file():
            return None, f"File not found, {USER_NAME}."

        try:
            ext = path.suffix.lower()
            if ext == ".pdf":
                import fitz

                doc = fitz.open(str(path))
                text = "".join(page.get_text() for page in doc)
            elif ext in (".docx", ".doc"):
                try:
                    from docx import Document

                    doc = Document(str(path))
                    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                except ImportError:
                    text = self._extract_text_from_docx_archive(path)
            else:
                text = path.read_text(encoding="utf-8", errors="ignore")
            return path, text[:max_chars]
        except Exception as exc:
            return None, f"Couldn't read file: {exc}"

    def find_file(self, name: str) -> str:
        if not name:
            return f"What file, {USER_NAME}?"
        results = []
        for folder in list(FOLDERS.values()) + [Path.home()]:
            if not folder.exists():
                continue
            for found in folder.rglob(f"*{name}*"):
                results.append(str(found))
                if len(results) >= 5:
                    break
            if len(results) >= 5:
                break
        if not results:
            return f"No files matching '{name}' found, {USER_NAME}."
        return "Found:\n" + "\n".join(f"- {item}" for item in results)

    def open_folder(self, folder: str) -> str:
        path = FOLDERS.get(folder.lower(), Path.home())
        if not path.exists():
            return f"Folder '{folder}' not found."
        try:
            if IS_WIN:
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(path)])
            else:
                subprocess.Popen(["xdg-open", str(path)])
            return f"Opened {folder} folder, {USER_NAME}."
        except Exception as exc:
            return f"Couldn't open folder: {exc}"

    def read_file(self, path_or_query: str) -> str:
        _, content = self.extract_text(path_or_query, max_chars=1500)
        return content

    def open_file(self, path_or_query: str) -> str:
        path = self._resolve_path(path_or_query)
        if not path.exists() or not path.is_file():
            return f"File not found, {USER_NAME}."
        return self._open_path(path, request=path_or_query)

    def open_recent_file(self, query: str = "") -> str:
        lookup_query = self._sanitized_file_query(query)
        path = self._find_recent_file(lookup_query or query)
        if not path and lookup_query and not self._is_recent_reference_query(query):
            resolved = self._resolve_path(lookup_query)
            if resolved.exists() and resolved.is_file():
                path = resolved

        if not path:
            return f"I couldn't find the recent file you meant, {USER_NAME}."
        return self._open_path(path, request=query or "open recent file")

    def create_file_interactive(self, description: str, ai_client) -> str:
        request = (description or "").strip()
        if not request:
            return f"Tell me what to create, {USER_NAME}."

        file_type = self._detect_file_type(request)
        document_kind = self._detect_document_kind(request)
        word_count = self._extract_word_count(request)
        prompt = self._build_generation_prompt(request, document_kind, file_type, word_count)
        content = ai_client.chat(prompt).strip()
        content = self._clean_generated_content(content)
        if not content:
            return f"I couldn't generate the file content yet, {USER_NAME}."

        title = self._derive_title(request, content, document_kind=document_kind)
        extension = ".docx" if file_type == "docx" else ".txt"
        label = "Word file" if extension == ".docx" else "text file"
        filename = self._build_filename(title, extension)
        path = self._next_available_path(self.output_dir / filename)

        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            if extension == ".docx":
                self._write_docx(path, title, content)
            else:
                path.write_text(content + ("\n" if not content.endswith("\n") else ""), encoding="utf-8")
        except Exception as exc:
            log.error(f"File creation failed for '{path}': {exc}")
            return f"I couldn't create that file yet, {USER_NAME}: {exc}"

        if not path.exists() or path.stat().st_size == 0:
            return f"I couldn't confirm that the file was saved, {USER_NAME}."

        self._remember_recent_file(path, request=request, title=title, kind=document_kind)
        return f"{label.capitalize()} created: {path}"

    def organize_by_type(self, folder: str = "downloads") -> str:
        src = FOLDERS.get(folder, FOLDERS["downloads"])
        moved = 0
        for file in src.iterdir():
            if file.is_dir():
                continue
            for dest_name, exts in TYPE_MAP.items():
                if file.suffix.lower() in exts:
                    dest = src / dest_name
                    dest.mkdir(exist_ok=True)
                    file.rename(dest / file.name)
                    moved += 1
                    break
        return f"Organized {moved} files in {folder}, {USER_NAME}."

    @staticmethod
    def _detect_file_type(description: str) -> str:
        low = description.lower()
        if any(term in low for term in ("word file", ".docx", "docx", "word document", "essay", "report", "resume", "cv", "biodata")):
            return "docx"
        if any(term in low for term in (".txt", "text file", "txt file")):
            return "txt"
        return "docx"

    @staticmethod
    def _detect_document_kind(description: str) -> str:
        low = description.lower()
        if any(term in low for term in ("resume", "cv", "curriculum vitae", "biodata", "bio data")):
            return "resume"
        if "essay" in low:
            return "essay"
        if any(term in low for term in ("cover letter", "application letter")):
            return "cover_letter"
        return "general"

    @staticmethod
    def _extract_word_count(description: str) -> int | None:
        match = re.search(r"\b(\d{2,4})\s*words?\b", description, re.IGNORECASE)
        if match:
            return int(match.group(1))
        return None

    def _build_generation_prompt(
        self,
        description: str,
        document_kind: str,
        file_type: str,
        word_count: int | None,
    ) -> str:
        if document_kind == "resume":
            return self._resume_prompt(description)
        prompt = self._content_prompt(description, file_type, word_count)
        memory_context = self._relevant_memory_context(description, max_items=6)
        if memory_context:
            prompt += (
                "\nRelevant memory you may use only if it directly helps this document:\n"
                f"{memory_context}\n"
            )
        return prompt

    def _resume_prompt(self, description: str) -> str:
        profile_context = self._profile_context(limit=18)
        recent_history = self._recent_history_context(limit=6)
        return (
            "Create a professional resume as plain text for a Word document.\n"
            f"User request: {description}\n"
            "Rules:\n"
            "- Use the stored profile details and recent conversation cues below when relevant.\n"
            "- Never invent personal details such as phone number, email, education, experience, or skills.\n"
            "- If information is missing, leave a clear placeholder like [Email], [Phone Number], or [Add Experience].\n"
            "- Prefer this structure: Full Name, Contact Information, Professional Summary, Skills, Experience, Education, Projects, Additional Details.\n"
            "- Keep the tone professional and ready to edit in Word.\n"
            "- Return only the resume text, with simple headings and body content.\n"
            f"Stored profile:\n{profile_context or 'No stored profile details yet.'}\n"
            f"Recent conversation:\n{recent_history or 'No recent personal context available.'}\n"
        )

    @staticmethod
    def _content_prompt(description: str, file_type: str, word_count: int | None) -> str:
        length_line = (
            f"- Keep the body close to {word_count} words.\n"
            if word_count
            else "- Keep the body concise and appropriate for the request.\n"
        )
        format_line = "Word document" if file_type == "docx" else "text file"
        return (
            "Write the exact content for this file request.\n"
            f"User request: {description}\n"
            f"Target format: {format_line}\n"
            "Rules:\n"
            f"{length_line}"
            "- Return only the document body text.\n"
            "- Do not add markdown, code fences, or explanations.\n"
            "- Do not say that the file was created or saved.\n"
            "- If the user asked for an essay, write the full essay directly.\n"
        )

    @staticmethod
    def _clean_generated_content(content: str) -> str:
        cleaned = (content or "").strip()
        if cleaned.startswith("```") and cleaned.endswith("```"):
            cleaned = "\n".join(cleaned.splitlines()[1:-1]).strip()
        return cleaned.strip().strip('"').strip()

    def _derive_title(self, description: str, content: str, document_kind: str = "general") -> str:
        if document_kind == "resume":
            resume_title = self._resume_title()
            if resume_title:
                return resume_title
            return "Resume"

        topic = self._extract_topic(description)
        if topic:
            if "essay" in description.lower():
                return f"{topic} Essay"
            return topic

        first_line = next((line.strip() for line in content.splitlines() if line.strip()), "")
        if first_line:
            words = first_line.split()
            return " ".join(words[:6]).strip(" ,.-") or "Jarvis Document"
        return "Jarvis Document"

    @staticmethod
    def _extract_topic(description: str) -> str:
        patterns = [
            r"\b(?:essay|report|article|document|file)\s+(?:on|about|for)\s+(.+)",
            r"\b(?:on|about)\s+(.+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE)
            if not match:
                continue
            topic = match.group(1).strip(" .,:-")
            topic = re.sub(r"\bof\s+\d+\s+words?\b.*$", "", topic, flags=re.IGNORECASE).strip(" .,:-")
            topic = re.sub(r"\b(?:in|on|as)\s+(?:a\s+)?(?:word\s+file|word\s+document|docx)\b.*$", "", topic, flags=re.IGNORECASE).strip(" .,:-")
            topic = re.sub(r"\b(?:please|do it|make it|save it|now)\b.*$", "", topic, flags=re.IGNORECASE).strip(" .,:-")
            topic = re.sub(r"\bfast\b.*$", "", topic, flags=re.IGNORECASE).strip(" .,:-")
            if topic:
                return " ".join(word.capitalize() for word in topic.split())
        return ""

    @staticmethod
    def _build_filename(title: str, extension: str) -> str:
        slug = re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_")
        slug = slug[:48] or "jarvis_document"
        return f"{slug}{extension}"

    def _resume_title(self) -> str:
        if not self.memory or not hasattr(self.memory, "get_profile_snapshot"):
            return ""
        snapshot = self.memory.get_profile_snapshot()
        name = snapshot.get("identity", {}).get("full_name", "").strip()
        if name:
            return f"{name} Resume"
        return ""

    def _profile_context(self, limit: int = 18) -> str:
        if not self.memory or not hasattr(self.memory, "get_profile_entries"):
            return ""
        entries = self.memory.get_profile_entries(limit=limit)
        if not entries:
            return ""
        lines = []
        for item in entries:
            lines.append(f"- {item['value']}")
        return "\n".join(lines)

    def _recent_history_context(self, limit: int = 6) -> str:
        if not self.memory or not hasattr(self.memory, "get_history"):
            return ""
        history = self.memory.get_history(limit=limit)
        if not history:
            return ""
        lines = []
        for item in history[-limit:]:
            user_text = (item.get("user") or "").strip()
            if user_text:
                lines.append(f"- User: {user_text}")
        return "\n".join(lines)

    def _relevant_memory_context(self, query: str, max_items: int = 6) -> str:
        if not self.memory or not hasattr(self.memory, "build_context"):
            return ""
        return self.memory.build_context(query, max_items=max_items)

    def _remember_recent_file(self, path: Path, request: str = "", title: str = "", kind: str = ""):
        record = {
            "path": Path(path),
            "request": (request or "").strip(),
            "title": (title or path.stem).strip(),
            "kind": (kind or "").strip().lower(),
        }
        self._recent_files = [
            item for item in self._recent_files
            if str(item["path"]).lower() != str(record["path"]).lower()
        ]
        self._recent_files.append(record)
        self._recent_files = self._recent_files[-12:]

    def _find_recent_file(self, query: str = "") -> Path | None:
        if not self._recent_files:
            return None

        cleaned = self._normalize_lookup(query)
        if not cleaned or self._is_recent_reference_query(query):
            return self._recent_files[-1]["path"]

        tokens = self._lookup_tokens(query)
        best = None
        best_score = 0
        for index, item in enumerate(self._recent_files):
            haystack = " ".join(
                [
                    item["title"],
                    item["request"],
                    item["path"].stem,
                    item["kind"],
                ]
            ).lower()
            score = sum(3 for token in tokens if token in haystack)
            if item["kind"] and item["kind"] in cleaned:
                score += 2
            score += index
            if score > best_score:
                best_score = score
                best = item["path"]

        if best_score > 0:
            return best
        return self._recent_files[-1]["path"] if self._is_recent_reference_query(query) else None

    @staticmethod
    def _normalize_lookup(text: str) -> str:
        return " ".join(re.findall(r"[a-z0-9]+", (text or "").lower()))

    def _lookup_tokens(self, text: str) -> list[str]:
        stop_words = {
            "open", "it", "that", "the", "file", "you", "created", "create",
            "for", "me", "please", "can", "jarvis", "document",
        }
        tokens = []
        for token in self._normalize_lookup(text).split():
            if token not in stop_words and len(token) > 1:
                tokens.append(token)
        return tokens

    def _is_recent_reference_query(self, text: str) -> bool:
        low = (text or "").lower().strip()
        if not low:
            return True
        phrases = [
            "open it",
            "open that",
            "open the file",
            "that file",
            "the file you created",
            "the file you create",
            "open recent file",
            "last file",
        ]
        if any(phrase in low for phrase in phrases):
            return True
        normalized = self._normalize_lookup(low)
        return normalized in {"it", "that", "that file", "open it", "open that"}

    @staticmethod
    def _sanitized_file_query(text: str) -> str:
        cleaned = (text or "").strip()
        cleaned = re.sub(r"^(?:please\s+)?open\s+", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"\b(?:the\s+)?file\s+you\s+(?:created|create)\b", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"\bit\b$", "", cleaned, flags=re.IGNORECASE).strip(" ,.-")
        return cleaned

    def _open_path(self, path: Path, request: str = "") -> str:
        try:
            if IS_WIN:
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(path)])
            else:
                subprocess.Popen(["xdg-open", str(path)])
            self._remember_recent_file(path, request=request, title=path.stem, kind=path.suffix.lstrip("."))
            return f"Opened file: {path}"
        except Exception as exc:
            return f"Couldn't open file: {exc}"

    @staticmethod
    def _next_available_path(path: Path) -> Path:
        if not path.exists():
            return path
        for index in range(2, 1000):
            candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
            if not candidate.exists():
                return candidate
        raise RuntimeError("Too many files with the same name already exist.")

    @staticmethod
    def _write_docx(path: Path, title: str, content: str):
        blocks = [block.strip() for block in re.split(r"\n\s*\n", content) if block.strip()]
        first_block = blocks[0] if blocks else ""
        normalized_title = re.sub(r"[^a-z0-9]+", "", title.lower())
        normalized_first = re.sub(r"[^a-z0-9]+", "", first_block.lower())

        if not blocks:
            blocks = [content.strip()]

        try:
            from docx import Document

            document = Document()
            if title and normalized_title and normalized_title != normalized_first:
                document.add_heading(title, level=0)
            for block in blocks:
                if block:
                    document.add_paragraph(block)
            document.save(str(path))
            return
        except ImportError:
            pass

        if title and normalized_title and normalized_title != normalized_first:
            blocks = [title] + blocks

        paragraphs = []
        for block in blocks:
            text = escape(block)
            text = text.replace("\n", "</w:t><w:br/><w:t xml:space=\"preserve\">")
            paragraphs.append(
                "<w:p><w:r><w:t xml:space=\"preserve\">"
                f"{text}"
                "</w:t></w:r></w:p>"
            )

        document_xml = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:document xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
            "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
            "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
            "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
            "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
            "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
            "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
            "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
            "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
            "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
            "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
            "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
            "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
            "xmlns:wne=\"http://schemas.microsoft.com/office/2006/wordml\" "
            "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" "
            "mc:Ignorable=\"w14 wp14\">"
            "<w:body>"
            + "".join(paragraphs)
            + "<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/>"
            "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" "
            "w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/>"
            "<w:cols w:space=\"708\"/><w:docGrid w:linePitch=\"360\"/></w:sectPr>"
            "</w:body></w:document>"
        )

        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr(
                "[Content_Types].xml",
                "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
                "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
                "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
                "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
                "<Override PartName=\"/word/document.xml\" "
                "ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
                "</Types>",
            )
            archive.writestr(
                "_rels/.rels",
                "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
                "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
                "<Relationship Id=\"rId1\" "
                "Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" "
                "Target=\"word/document.xml\"/>"
                "</Relationships>",
            )
            archive.writestr("word/document.xml", document_xml)

    @staticmethod
    def _extract_text_from_docx_archive(path: Path) -> str:
        with zipfile.ZipFile(path, "r") as archive:
            xml_text = archive.read("word/document.xml").decode("utf-8", errors="ignore")

        xml_text = re.sub(r"</w:p>", "\n", xml_text)
        xml_text = re.sub(r"<w:br\s*/>", "\n", xml_text)
        xml_text = re.sub(r"<[^>]+>", "", xml_text)
        xml_text = xml_text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        return "\n".join(line.strip() for line in xml_text.splitlines() if line.strip())
